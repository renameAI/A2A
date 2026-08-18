"""업로드 자료: 스토리지 경유 + Word(.docx) 파싱.

배경(실측): 엔진이 업로드 파일을 로컬 디스크에 쓰고 나중에 경로로 읽었는데,
Vercel 서버리스는 /tmp 외 읽기 전용이고 /tmp조차 호출 간 공유되지 않는다.
193바이트 PDF도 500이 났다. 그리고 파일이 Vercel 함수를 통과하면 요청 본문
4.5MB 상한(413)에 걸린다 — 요금제로 올릴 수 없는 플랫폼 제한이다.
"""
import io

import pytest

from app.ingest.chunking import docx_to_text
from app.saas import storage


def _docx(paragraphs=(), table=None) -> bytes:
    from docx import Document
    d = Document()
    for t in paragraphs:
        d.add_paragraph(t)
    if table:
        tb = d.add_table(rows=len(table), cols=len(table[0]))
        for i, row in enumerate(table):
            for j, cell in enumerate(row):
                tb.cell(i, j).text = cell
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


class TestDocx:
    def test_paragraphs_and_tables_are_both_extracted(self):
        """회사 소개 문서는 실적·고객사를 표에 담는다 — 본문만 뽑으면 사라진다."""
        out = docx_to_text(_docx(
            paragraphs=["뉴톤은 MRV를 디지털화합니다."],
            table=[["적용 분야", "바이오차"], ["실증 지역", "베트남"]]))
        assert "뉴톤은 MRV를 디지털화합니다." in out
        assert "적용 분야 | 바이오차" in out
        assert "베트남" in out

    def test_blank_paragraphs_are_dropped(self):
        out = docx_to_text(_docx(paragraphs=["A", "   ", "", "B"]))
        assert out.split("\n") == ["A", "B"]

    def test_merged_cells_are_not_repeated(self):
        """병합 셀은 같은 텍스트를 여러 번 돌려준다 — 한 번만 남아야 한다."""
        out = docx_to_text(_docx(table=[["같음", "같음"]]))
        assert out == "같음"

    def test_not_a_docx_raises_rather_than_returning_garbage(self):
        with pytest.raises(Exception):
            docx_to_text(b"%PDF-1.4\nnot a docx")


class TestStorageScheme:
    def test_supabase_scheme_reads_through_storage(self, monkeypatch):
        """supabase:// 자산은 로컬 경로가 아니라 스토리지에서 읽힌다."""
        from app.ingest import fetchers
        seen = {}

        def _download(obj):
            seen["obj"] = obj
            return b"%PDF-1.4\n"
        monkeypatch.setattr(storage, "download", _download)
        data = fetchers.fetch_pdf_bytes("supabase://ws-boram/abc.pdf", None)
        assert seen["obj"] == "ws-boram/abc.pdf"
        assert data.startswith(b"%PDF-")

    def test_missing_object_becomes_fetch_failed(self, monkeypatch):
        """스토리지에 없는 자산은 FetchFailed로 떨어져 상위가 원인을 안다."""
        from app.ingest import fetchers
        from app.ingest.fetchers import FetchFailed

        def _boom(obj):
            raise RuntimeError("404")
        monkeypatch.setattr(storage, "download", _boom)
        with pytest.raises(FetchFailed):
            fetchers.fetch_pdf_bytes("supabase://ws-boram/gone.pdf", None)


class TestContentTypes:
    def test_bucket_and_engine_agree_on_accepted_types(self):
        """엔진이 통과시킨 형식을 스토리지가 거절하면 사용자는 이유를 모른다."""
        assert set(storage.CONTENT_TYPES) == {".pdf", ".docx"}
        assert storage.CONTENT_TYPES[".pdf"] == "application/pdf"
        assert storage.CONTENT_TYPES[".docx"].endswith("wordprocessingml.document")
