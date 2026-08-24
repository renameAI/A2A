"""청킹 — 출처 라벨을 유지하는 분할 (ING-02).

PDF는 페이지 단위로 텍스트를 뽑은 뒤 문단 경계 우선으로 ~2,000자 청크를 만든다.
청크 ID는 provenance 역추적(ING-04)의 키가 된다.
"""
import io
from dataclasses import dataclass


DEFAULT_CHUNK_CHARS = 2000


@dataclass
class Chunk:
    chunk_id: str    # 예: "a1:ir_deck#3"
    source: str      # 자산 라벨 (예: "a1:ir_deck")
    text: str


def pdf_to_text(data: bytes) -> str:
    # 지연 임포트 — pypdf는 cryptography까지 끌고 와 콜드스타트에 ~160ms를
    # 더한다(실측 importtime). PDF 인제스트 때만 필요한 비용을 모든 요청이
    # 내고 있었다. 아래 docx와 같은 패턴.
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    pages = []
    for i, page in enumerate(reader.pages):
        text = (page.extract_text() or "").strip()
        if text:
            pages.append(f"[p.{i + 1}]\n{text}")
    return "\n\n".join(pages)


def docx_to_text(data: bytes) -> str:
    """Word(.docx) → 텍스트.

    LLM에 맡기지 않는 이유: .docx는 zip 안의 XML이라 텍스트가 **이미 구조로
    존재한다**. 모델을 통과시키면 비용이 들고, 없던 문장이 생길 여지가 생기며,
    결과가 실행마다 흔들린다. 읽어낼 수 있는 것은 읽어낸다.

    표를 함께 뽑는 이유: 회사 소개 문서는 실적·고객사·스펙을 표에 담는 경우가
    많다. 본문만 뽑으면 그 부분이 통째로 사라진다.
    """
    from docx import Document
    doc = Document(io.BytesIO(data))
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for t in doc.tables:
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                # 중복 셀 제거 — 병합 셀은 같은 텍스트를 여러 번 돌려준다
                seen, uniq = set(), []
                for c in cells:
                    if c not in seen:
                        seen.add(c); uniq.append(c)
                parts.append(" | ".join(uniq))
    return "\n".join(parts)


def chunk_text(text: str, source: str,
               max_chars: int = DEFAULT_CHUNK_CHARS) -> list[Chunk]:
    """문단 경계 우선 분할. 문단이 max_chars를 넘으면 그 안에서 하드 분할."""
    paragraphs: list[str] = []
    for para in text.split("\n\n"):
        para = para.strip()
        if not para:
            continue
        while len(para) > max_chars:
            paragraphs.append(para[:max_chars])
            para = para[max_chars:]
        paragraphs.append(para)

    chunks: list[Chunk] = []
    buffer = ""
    for para in paragraphs:
        if buffer and len(buffer) + len(para) + 2 > max_chars:
            chunks.append(Chunk(f"{source}#{len(chunks) + 1}", source, buffer))
            buffer = para
        else:
            buffer = f"{buffer}\n\n{para}" if buffer else para
    if buffer:
        chunks.append(Chunk(f"{source}#{len(chunks) + 1}", source, buffer))
    return chunks
